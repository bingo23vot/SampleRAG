
import streamlit as st
import os
import pickle
from typing import List, Dict
from pathlib import Path
import tempfile

# Document processing imports
import PyPDF2
from docx import Document
import pandas as pd

# LangChain imports
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_core.documents import Document as LangChainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain_core.prompts import PromptTemplate


def load_config() -> Dict[str, str]:
    """Load configuration from application.properties file"""
    config = {}
    try:
        with open('application.properties', 'r') as f:
            for line in f:
                if '=' in line and not line.strip().startswith('#'):
                    key, value = line.strip().split('=', 1)
                    config[key.strip()] = value.strip().strip('"')
    except FileNotFoundError:
        st.error("application.properties file not found")
    except Exception as e:
        st.error(f"Error loading config: {str(e)}")
    return config


config = load_config()
if 'GEMINI_KEY' in config:
    os.environ["GOOGLE_API_KEY"] = config['GEMINI_KEY']


class DocumentProcessor:
    """Class to handle document processing for various file types"""

    def __init__(self) -> None:
        self.supported_types = ['.pdf', '.docx', '.doc', '.xlsx', '.xls']

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except FileNotFoundError:
            st.error(f"PDF file not found: {file_path}")
        except PyPDF2.errors.PdfReadError:
            st.error("Error reading PDF: File may be corrupted or encrypted")
        except Exception as e:
            st.error(f"Error reading PDF: {str(e)}")
        return text

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except FileNotFoundError:
            st.error(f"DOCX file not found: {file_path}")
        except Exception as e:
            st.error(f"Error reading DOCX: {str(e)}")
        return text

    def extract_text_from_excel(self, file_path: str) -> str:
        """Extract text from Excel file"""
        text = ""
        try:
            excel_file = pd.ExcelFile(file_path)
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                text += f"Sheet: {sheet_name}\n"
                text += df.to_string() + "\n\n"
        except FileNotFoundError:
            st.error(f"Excel file not found: {file_path}")
        except Exception as e:
            st.error(f"Error reading Excel: {str(e)}")
        return text

    def process_document(self, file_path: str, file_extension: str) -> str:
        """Process document based on file type"""
        file_ext_lower = file_extension.lower()
        
        if file_ext_lower == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext_lower in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        elif file_ext_lower in ['.xlsx', '.xls']:
            return self.extract_text_from_excel(file_path)
        else:
            st.error(f"Unsupported file type: {file_extension}")
            return ""


class RAGSystem:
    """RAG (Retrieval-Augmented Generation) system using LangChain and Gemini"""

    def __init__(self) -> None:
        # Initialize LangChain components
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001")
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-2.0-flash-exp",
            temperature=0)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", " ", ""])

        self.doc_processor = DocumentProcessor()
        self.vectorstore = None
        self.documents = []
        self.vectorstore_file = "vectorstore"
        self.documents_file = "documents.pkl"
        self.load_existing_data()

    def load_existing_data(self) -> None:
        """Load existing vectorstore and documents"""
        try:
            if os.path.exists(self.documents_file):
                with open(self.documents_file, 'rb') as f:
                    self.documents = pickle.load(f)

            if os.path.exists(f"{self.vectorstore_file}.faiss"):
                self.vectorstore = FAISS.load_local(
                    self.vectorstore_file,
                    self.embeddings,
                    allow_dangerous_deserialization=True)
        except FileNotFoundError:
            st.warning("No existing data files found. Starting fresh.")
        except Exception as e:
            st.error(f"Error loading existing data: {str(e)}")
            self.documents = []
            self.vectorstore = None

    def save_data(self) -> None:
        """Save vectorstore and documents to files"""
        try:
            if self.vectorstore:
                self.vectorstore.save_local(self.vectorstore_file)

            with open(self.documents_file, 'wb') as f:
                pickle.dump(self.documents, f)
        except Exception as e:
            st.error(f"Error saving data: {str(e)}")

    def add_document(self, file_name: str, text: str) -> int:
        """Add document using LangChain"""
        if text.strip():
            # Create LangChain document
            langchain_doc = LangChainDocument(
                page_content=text,
                metadata={"source": file_name})

            # Split into chunks
            chunks = self.text_splitter.split_documents([langchain_doc])

            # Store document info
            doc_info = {
                'file_name': file_name,
                'text': text,
                'chunks': len(chunks)
            }
            self.documents.append(doc_info)

            # Create or update vectorstore
            if self.vectorstore is None:
                self.vectorstore = FAISS.from_documents(
                    chunks, self.embeddings)
            else:
                self.vectorstore.add_documents(chunks)

            self.save_data()
            return len(chunks)
        return 0

    def search_and_answer(self, query: str, top_k: int = 5) -> Dict[str, List]:
        """Search documents and generate answer using LangChain RAG"""
        if not self.vectorstore:
            return {
                "answer": "No documents available for search.",
                "sources": []
            }

        try:
            # Create retrieval QA chain
            prompt_template = """Use the following pieces of context to answer the question at the end. 
            If you don't know the answer, just say that you don't know, don't try to make up an answer.

            Context: {context}

            Question: {question}
            Answer:"""

            prompt = PromptTemplate(
                template=prompt_template,
                input_variables=["context", "question"])

            qa_chain = RetrievalQA.from_chain_type(
                llm=self.llm,
                chain_type="stuff",
                retriever=self.vectorstore.as_retriever(
                    search_kwargs={"k": top_k}),
                chain_type_kwargs={"prompt": prompt},
                return_source_documents=True)

            # Get answer
            result = qa_chain({"query": query})

            # Extract source information
            sources = []
            for langchain_doc in result.get("source_documents", []):
                content = langchain_doc.page_content
                truncated_content = (content[:300] + "..." 
                                   if len(content) > 300 else content)
                sources.append({
                    "content": truncated_content,
                    "source": langchain_doc.metadata.get("source", "Unknown"),
                    "full_content": content
                })

            return {"answer": result["result"], "sources": sources}

        except Exception as e:
            st.error(f"Error in search and answer: {str(e)}")
            return {
                "answer": "An error occurred while processing your query.",
                "sources": []
            }

    def search_similar_documents(self, query: str, top_k: int = 5) -> List[Dict]:
        """Search for similar documents using vectorstore"""
        if not self.vectorstore:
            return []

        try:
            docs = self.vectorstore.similarity_search_with_score(query, k=top_k)
            results = []

            for langchain_doc, score in docs:
                result = {
                    'file_name': langchain_doc.metadata.get("source", "Unknown"),
                    'text': langchain_doc.page_content,
                    'similarity': 1 - score,  # Convert distance to similarity
                    'full_content': langchain_doc.page_content
                }
                results.append(result)

            return results
        except Exception as e:
            st.error(f"Error searching documents: {str(e)}")
            return []


def main() -> None:
    """Main application function"""
    st.set_page_config(
        page_title="RAG Document Upload System",
        page_icon="📚",
        layout="wide")

    st.title("📚 RAG Document Upload & Search System (LangChain)")
    st.markdown(
        "Upload documents (PDF, DOC, DOCX, XLS, XLSX) and search through them using LangChain with Gemini embeddings"
    )

    # Initialize RAG system
    if 'rag_system' not in st.session_state:
        st.session_state.rag_system = RAGSystem()

    rag_system = st.session_state.rag_system

    # Sidebar for document upload
    with st.sidebar:
        st.header("📤 Upload Documents")

        uploaded_files = st.file_uploader(
            "Choose files",
            type=['pdf', 'docx', 'doc', 'xlsx', 'xls'],
            accept_multiple_files=True)

        if uploaded_files:
            if st.button("Process Documents"):
                progress_bar = st.progress(0)
                total_files = len(uploaded_files)

                for idx, uploaded_file in enumerate(uploaded_files):
                    st.write(f"Processing: {uploaded_file.name}")

                    # Save uploaded file temporarily
                    with tempfile.NamedTemporaryFile(
                            delete=False, 
                            suffix=Path(uploaded_file.name).suffix) as tmp_file:
                        tmp_file.write(uploaded_file.getvalue())
                        tmp_file_path = tmp_file.name

                    try:
                        # Extract text
                        file_extension = Path(uploaded_file.name).suffix
                        text = rag_system.doc_processor.process_document(
                            tmp_file_path, file_extension)

                        if text:
                            chunks_added = rag_system.add_document(
                                uploaded_file.name, text)
                            st.success(
                                f"Added {chunks_added} chunks from {uploaded_file.name}"
                            )
                        else:
                            st.error(
                                f"Could not extract text from {uploaded_file.name}"
                            )

                    except Exception as e:
                        st.error(
                            f"Error processing {uploaded_file.name}: {str(e)}")

                    finally:
                        # Clean up temporary file
                        if os.path.exists(tmp_file_path):
                            os.unlink(tmp_file_path)

                    progress_bar.progress((idx + 1) / total_files)

                st.success("All documents processed!")
                st.rerun()

    # Main area for search and results
    col1, col2 = st.columns([1, 2])

    with col1:
        st.header("🔍 Search Documents")

        # Display current document count
        st.info(f"📊 Documents in database: {len(rag_system.documents)}")
        total_chunks = sum(
            doc.get('chunks', 0) for doc in rag_system.documents)
        st.info(f"📝 Total chunks: {total_chunks}")

        # Search interface
        search_query = st.text_area("Enter your search query:", height=100)
        top_k = st.slider("Number of results:",
                          min_value=1,
                          max_value=10,
                          value=5)

        search_type = st.radio("Search Type:",
                               ["Answer with AI", "Similar Documents"])

        if st.button("Search") and search_query:
            with st.spinner("Searching..."):
                if search_type == "Answer with AI":
                    result = rag_system.search_and_answer(search_query, top_k)
                    st.session_state.search_result = result
                    st.session_state.search_type = "answer"
                else:
                    results = rag_system.search_similar_documents(
                        search_query, top_k)
                    st.session_state.search_results = results
                    st.session_state.search_type = "similarity"

    with col2:
        st.header("📋 Search Results")

        if hasattr(st.session_state, 'search_type'):
            if (st.session_state.search_type == "answer" and 
                hasattr(st.session_state, 'search_result')):
                result = st.session_state.search_result

                st.subheader("🤖 AI Answer:")
                st.write(result["answer"])

                if result["sources"]:
                    st.subheader("📚 Sources:")
                    for i, source in enumerate(result["sources"]):
                        with st.expander(f"Source {i+1}: {source['source']}"):
                            st.write(source["content"])
                            if st.button(f"Show full content",
                                         key=f"show_source_{i}"):
                                st.text_area("Full content:",
                                             source["full_content"],
                                             height=300,
                                             key=f"source_full_{i}")

            elif (st.session_state.search_type == "similarity" and 
                  hasattr(st.session_state, 'search_results')):
                if st.session_state.search_results:
                    for i, result in enumerate(st.session_state.search_results):
                        similarity_score = result['similarity']
                        with st.expander(
                                f"Result {i+1}: {result['file_name']} "
                                f"(Similarity: {similarity_score:.3f})"):
                            st.write("**Text snippet:**")
                            st.write(result['text'])

                            if st.button(f"Show full content",
                                         key=f"show_full_{i}"):
                                st.text_area("Full content:",
                                             result['full_content'],
                                             height=300,
                                             key=f"full_text_{i}")
                else:
                    st.info("No results found. Try a different query.")
        else:
            st.info("Enter a search query to see results here.")

    # Management section
    st.header("🗂️ Database Management")
    col1, col2 = st.columns(2)

    with col1:
        if st.button("Clear All Documents"):
            rag_system.documents = []
            rag_system.vectorstore = None
            # Clean up saved files
            faiss_file = f"{rag_system.vectorstore_file}.faiss"
            pkl_file = f"{rag_system.vectorstore_file}.pkl"
            
            if os.path.exists(faiss_file):
                os.remove(faiss_file)
            if os.path.exists(pkl_file):
                os.remove(pkl_file)
            if os.path.exists(rag_system.documents_file):
                os.remove(rag_system.documents_file)
                
            rag_system.save_data()
            st.success("All documents cleared!")
            st.rerun()

    with col2:
        if rag_system.documents:
            st.download_button(
                label="Download Document Database",
                data=pickle.dumps(rag_system.documents),
                file_name="document_database.pkl",
                mime="application/octet-stream")


if __name__ == "__main__":
    main()
